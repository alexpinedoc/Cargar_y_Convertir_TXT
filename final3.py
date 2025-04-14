"""
Nombre: Cargar_txt_a_BD_y_convertir_a_pdf.py
Version: 1.0
Autor: Alex Pinedo
Fecha: 14/04/2025
Descripcion: Este script tiene 3 funciones:
    1. Procesar los archivos .txt, extraer datos específicos para ser usados en una plantilla 
        de Word. Luego, convierte el documento a PDF, lo guarda (junto a su txt) en una carpeta de destino.
    2. Realizar una inserción a la tabla "CTABNCO_DOCVIRTUAL" de la BD Oracle, usando los datos 
        extraídos de cada archivo .txt.
    3. (opcional) Si hay archivos .txt que no se pudieron procesar correctamente, se enviara
        un correo, adjuntando los archivos .txt y los errores correspondiente.

Parametros: 
    Las carpetas y plantilla:
    1. carpeta_origen: Ruta de la carpeta donde se encuentran los archivos .txt a procesar.
    2. carpeta_destino: Ruta de la carpeta donde se guardaran los archivos .pdf y .txt procesados.
    3. carpeta_destino_error: Ruta de la carpeta donde se guardaran los archivos .txt que no  
        se pudieron procesar correctamente.
    4. ruta_plantilla: Ruta de la plantilla de Word que se usara para crear los documentos .docx.

    Conexion a la BD Oracle:
    1. username: Nombre de usuario para la conexión a la BD Oracle.
    2. password: Contraseña para la conexión a la BD Oracle.
    3. dsn: Nombre del servicio de la BD Oracle, colocas la ip y el nombre.

    Configuración del envio de correos:
    1. smtp_server: Servidor SMTP de gmail.
    2. smtp_port: Puerto para SSL.
    3. email_origen: Email que enviara los correos.
    4. email_password: Contraseña de la aplicación del Email.
    5. destinatarios: Lista de Email's de destino.

    Otros:
    1. codificaciones: Lista de codificaciones a probar al abrir los archivos .txt.
    2. tipo: Tipo de archivo a insertar en la tabla "CTABNCO_DOCVIRTUAL". 

Notas:
    1. La plantilla de Word usa marcadores como (1), (2), etc. para indicar donde se
        deben insertar los datos extraídos del archivo .txt.
        1.1. Para la plantilla se usaron 12 de esos marcadores, de los cuales (3) y (9) son
            obtenidos de la BD Oracle.
    2. Si no hay errores en el proceso, no se enviara ningun correo.
    3. Se configuro "variables de entorno" en mi SO Windows, para la conexion a la BD Oracle 
        y la "clave de aplicacion" del Email.
"""

import os
import re
from datetime import datetime
import shutil #Para mover los archivos
import comtypes.client #Para convertir a PDF
from docx import Document #Para trabajar con Word
# Para el envio a correo
import smtplib
import ssl
from email.message import EmailMessage

import oracledb 
# Activa el modo "thick" 
oracledb.init_oracle_client(lib_dir=r"C:\Oracle\instantclient_23_7") 

# Lista de codificaciones a probar
codificaciones = ["utf-8", "latin-1", "cp1252"]

#CTABNCO_TXT
tipo = 3 

# Rutas de las carpetas y plantilla
carpeta_origen = r"\\25.0.3.9\archivos\archivo_temporal" 
carpeta_destino = r"\\25.0.3.9\archivos\archivo_finanzas"
carpeta_destino_error = r"\\25.0.3.9\archivos\archivo_temporal\carpeta_errores"
ruta_plantilla = r"\\25.0.3.9\archivos\archivo_temporal\Plantillas\ctabnco_txt\Plantilla_ctabnco_txt.docx"

# Configura la conexión a la BD
username = os.getenv("DB_USER")
password = os.getenv("DB_PASS")
dsn = oracledb.makedsn("10.0.0.100", 1521, sid="PEVISA10")

# Configuración del correo
smtp_server = "smtp.gmail.com"  # Servidor SMTP de Gmail
smtp_port = 465  # Puerto para SSL
email_origen = "gestion.errores.hdi@gmail.com"
email_password = os.getenv("EMAIL_CLAVE") # Usa la contraseña de aplicación generada

# Lista de destinatarios
destinatarios = ["apinedo@hdi.com.pe"]


# Almacenara los datos que obtengamos de los txt
resultados_1 = [] 

# Alamcenara los datos que tengan errores
resultados_error = []




##############################################################################################
"""SACAMOS LOS DATOS DEL TXT"""
def llenar_lista_resultados_temp (archivo, primera_linea, lista_resultados_temp):
    """
    Este proceso se encarga de llenar la lista de resultados temporales con los datos
    extraidos del archivo txt. La primera linea del archivo ha sido leido y esta como 
    parametro. 
    """
    # Leer las siguientes líneas del archivo   
    segunda_linea = archivo.readline()
    tercera_linea = archivo.readline()
    
    # En esta parte "archivo" esta en la linea 3, por lo que veremos cuantas
    # lineas tiene a partir de ahi (Si solo tiene 3 lineas, te dara 1)
    num_lineas = sum(1 for _ in archivo) + 1  

    # (1)     
    dato = primera_linea[213:219].strip()  # Extraer el código según la posición
    lista_resultados_temp.append(dato)

    # (2)
    dato = primera_linea[20:33].strip()
    lista_resultados_temp.append(dato)

    # (4)
    dato = primera_linea[57:80].strip()
    lista_resultados_temp.append(dato)

    # (5)
    dato = segunda_linea[154:171].strip()
    lista_resultados_temp.append(dato)

    # (6)
    dato = primera_linea[7:15].strip()
    dato_fecha = datetime.strptime(dato, "%Y%m%d").strftime("%d/%m/%Y")
    lista_resultados_temp.append(dato_fecha)

    # (7)
    dato = primera_linea[40:57].strip()
    dato_numero = str(float(dato))
    lista_resultados_temp.append(dato_numero)

    # (8)
    dato = segunda_linea[39:97].strip()
    lista_resultados_temp.append(dato)

    # (10)
    dato = tercera_linea[2:17].strip()
    lista_resultados_temp.append(dato)

    # (11)
    dato = datetime.now().strftime("%d/%m/%y, %H:%M")      
    lista_resultados_temp.append(dato)

    # (12)
    lista_resultados_temp.append(str(num_lineas))


def extraer_dato_pagos(primera_linea):
    """
    Esta funcion se encarga de extraer el indice que esta luego del dato "PAGOS" 
    de la primera linea del archivo.
    """
    match = re.search(r'PAGOS\s+(\S+)', primera_linea)  # Captura el valor después de "PAGOS"
    if match:

        return match.group(1)  # Devuelve el indice después de "PAGOS"
      
    return None


def obtener_datos(ruta_archivo, lista_resultados):
    """
    Esta proceso se encarga de obtener los datos del archivo txt y alamcenarlos
    en la tupla "lista_resultados".
    Primero, intenta abrir el archivo con las codificaciones especificadas.
    Segundo, lee la primera linea para que sea utilizada en la funcion "extraer_dato_pagos".
    Tercero, llama al proceso "llenar_lista_resultados_temp" para que llene la lista de resultados temporal 
    Cuarto, con los datos obtenidos de las funciones, se llena una fila de la tupla "lista_resultados" 
    """
    lista_resultados_temp = [] # Tupla de los datos dentro del txt actual
    dato_pagos = None # indice de pago (que usaremos para insertar a CTABNCO_DOCVIRTUAL)

    for encoding in codificaciones:
        try:
            with open(ruta_archivo, "r", encoding = encoding) as archivo:
                primera_linea = archivo.readline()  # Leer la primera línea
                dato_pagos = extraer_dato_pagos (primera_linea)
                llenar_lista_resultados_temp (archivo, primera_linea, lista_resultados_temp)  
        except UnicodeDecodeError:
            print(f"Fallo con {encoding}, intentando otra...")

    if dato_pagos:
        dato1, dato2 = dato_pagos.split("-")
        nombre_archivo = os.path.basename(ruta_archivo)
        nombre_archivo_pdf = os.path.splitext(nombre_archivo)[0] + ".pdf"  # Cambiar extensión
        lista_resultados.append((
            dato1, dato2, nombre_archivo, nombre_archivo_pdf,
            lista_resultados_temp[0], 
            lista_resultados_temp[1],
            lista_resultados_temp[2],
            lista_resultados_temp[3],
            lista_resultados_temp[4],
            lista_resultados_temp[5],
            lista_resultados_temp[6],
            lista_resultados_temp[7],
            lista_resultados_temp[8],
            lista_resultados_temp[9]
            ))
    else:
        print(f"No se encontró el dato 'PAGOS ...' en {ruta_archivo}")


def procesar_carpeta(carpeta_origen):
    """
    Esta funcion se encarga de procesar todos los archivos dentro de la carpeta especificada
    y llenar la tupla "lista_resultados" con los datos obtenidos de cada archivo.
    Primero, se recorre la carpeta y se verifica que cada elemento sea un archivo.
    Segundo, se llama a la funcion "obtener_datos" para que llene la tupla "lista_resultados".
    Tercero, se retorna la tupla "lista_resultados" con los datos obtenidos de todos los archivos.
    """
    lista_resultados = []  # Lista final de datos

    for archivo in os.listdir(carpeta_origen):
        ruta_archivo = os.path.join(carpeta_origen, archivo)
        if os.path.isfile(ruta_archivo):  # Verifica que sea un archivo y no una carpeta            

            obtener_datos(ruta_archivo, lista_resultados)

    return lista_resultados  # Retorna la lista de tuplas


#Obtenemos los valores de los txt y los almacenamos en la tupla "resultados_1"
resultados_1 = procesar_carpeta(carpeta_origen)
#print (resultados_1)



##############################################################################################
"""OBTENEMOS LOS DATOS (3) Y (9) DE LA BD"""
resultados_2 = [] # La idea es: mantener los 4 primeros valores como el indice y los nombres del archivo
                  # Lo demas, los valores que usaremos en el word - pdf

resultados_2_1 = [] # Almacena los datos que no tengan errores de "resultados_2"

try:
    connection = oracledb.connect(user=username, password=password, dsn=dsn)

    cursor_select = connection.cursor()
    cursor_insert = connection.cursor()

    sql_select = """SELECT t1.CIA, t2.ruc 
                    FROM pla_control t1
                    CROSS JOIN (SELECT distinct ruc 
                                FROM proveed 
                                WHERE nombre = :1) t2"""
    
    for row in resultados_1:

        dato_extraido = (row[10],) #Es la razón social / esa coma al final es para que python sepa que es una tupla
        cursor_select.execute(sql_select , dato_extraido)
        r_select = cursor_select.fetchone()
        resultados_2.append((row[0],row[1],row[2],row[3],
                            row[4],row[5],r_select[0],
                            row[6],row[7],row[8],row[9],
                            row[10],r_select[1],
                            row[11],row[12],row[13]))
        #El doble parentesis, es para que lo almacene como una tupla cada iteración

    #print(resultados_2);



    ##############################################################################################
    """HACEMOS LA INSERSION A CTABNCO_DOCVIRTUAL"""

    sql_select_2 = """SELECT DISTINCT voucher_caja, ano, mes, libro_caja 
                      FROM pagos_i 
                      WHERE serie_planilla = :1 
                      AND numero_planilla= :2"""
    
    sql_insert = """INSERT INTO CTABNCO_DOCVIRTUAL 
                        (id_tipo_archivo, codigo, ano, mes, libro, path_archivo, path_archivo_2) 
                    VALUES (:1, :2, :3, :4, :5, :6, :7)"""

    for row in resultados_2:
        datos_extraidos = (row[0], row[1]) # Es la serie y el numero de la planilla
        cursor_select.execute(sql_select_2 , datos_extraidos)
        r_select_2 = cursor_select.fetchone() 

        if (r_select_2 is not None):
            datos_insert = (tipo, r_select_2[0], r_select_2[1], r_select_2[2], r_select_2[3],
                            row[2], row[3])
            cursor_insert.execute(sql_insert, datos_insert)

            resultados_2_1.append(row)
        else:
            # Si no hay detalle en "pagos_i", se guarda el error
            resultados_error.append(row + ('Error: La planilla no tiene detalle, archivo: ' + row[2],))

    cursor_select.close()
    cursor_insert.close()
    connection.commit()
    connection.close()
    print("✅ Inserción completa")

except oracledb.Error as e:
    print("❌ Error:", e)

#print(resultados_2_1)
#print(resultados_error)



##############################################################################################
"""USAMOS LA LISTA Y LA PLANTILLA PARA CREAR LOS WORDS"""

for row in resultados_2_1:
    
    # Cargar la plantilla
    doc = Document(ruta_plantilla)

    reemplazos = {
        "(1)": row[4],
        "(2)": row[5],
        "(3)": row[6],
        "(4)": row[7],
        "(5)": row[8],
        "(6)": row[9],
        "(7)": row[10],
        "(8)": row[11],
        "(9)": row[12],
        "(10)": row[13],
        "(11)": row[14],
        "(12)": row[15]
    }

    
    def reemplazar_en_runs(parrafo):
        """Proceso para reemplazar texto sin perder formato en runs"""
        full_text = "".join(run.text for run in parrafo.runs)  # Unir el texto de todos los runs
        modificado = False
        for key, value in reemplazos.items():
            if key in full_text:
                full_text = full_text.replace(key, value)
                modificado = True
        if modificado:  # Si hubo cambios, actualizamos el párrafo
            for run in parrafo.runs:
                run.text = ""  # Borrar el contenido anterior
            parrafo.runs[0].text = full_text  # Escribir el nuevo texto en el primer run

    # Reemplazo en párrafos normales
    for parrafo in doc.paragraphs:
        
        reemplazar_en_runs(parrafo)

    # Reemplazo en tablas
    for table in doc.tables:
        for row_1 in table.rows:
            for cell in row_1.cells:
                for parrafo in cell.paragraphs:
                    reemplazar_en_runs(parrafo)
    
    """
    # Espacio de nombres XML de Word
    NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Reemplazo en encabezados y pies de página
    for section in doc.sections:
        # Encabezados
        for para in section.header.paragraphs:
            reemplazar_en_runs(parrafo)
        # Pies de página
        for para in section.footer.paragraphs:
            reemplazar_en_runs(parrafo)

        # Buscar texto dentro de figuras en el pie de página
        footer_xml = section.footer._element
        for shape in footer_xml.findall(".//w:t", namespaces=NS):  # Buscar los textos dentro de shapes
            text = shape.text
            for key, value in reemplazos.items():
                if text and key in text:
                    shape.text = text.replace(key, value)  # Reemplaza el texto dentro de la figura
    """

    # Ruta del documento modificado
    archivo = os.path.splitext(row[2])[0]
    new_doc_path = os.path.join(carpeta_origen, archivo + ".docx")
    doc.save(new_doc_path)


    ##############################################################################################
    """CONVERTIMOS LOS WORDS A PDF"""
    # Ruta donde ira el PDF
    pdf_path = os.path.join(carpeta_destino, archivo + ".pdf")

    # Convertir a PDF con Microsoft Word
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False  # No mostrar Word al ejecutar
    doc = word.Documents.Open(new_doc_path)
    doc.SaveAs(pdf_path, FileFormat=17)  # 17 = formato PDF
    doc.Close()
    word.Quit()

    # Verificar si el PDF se creó correctamente
    if os.path.exists(pdf_path):
        os.remove(new_doc_path)  # Eliminar el archivo .docx
        shutil.move(os.path.join(carpeta_origen, archivo + ".txt"), carpeta_destino) # Mover el txt a la carpeta de destino
        print(f"Documento convertido y guardado en: {pdf_path}")
        #print(f"Archivo .docx eliminado: {new_doc_path}")
    else:
        print("Error al convertir el archivo.")


##############################################################################################
"""CORREO DE LOS TXT QUE TUVIERON ERRORES"""

if resultados_error != []:
    # Crear el mensaje
    mensaje = EmailMessage()
    mensaje_error = ''

    #Genereamos el mensaje del correo
    for row in resultados_error:
        mensaje_error += f"Planilla {row[0]}-{row[1]}: {row[-1]}\n\n"       
        

    # Estructura del mensaje
    mensaje["Subject"] = "Aviso de error en las planillas"
    mensaje["From"] = email_origen
    mensaje["To"] = ", ".join(destinatarios)  # Unir los correos con comas
    mensaje.set_content(
"""Buen dia estimados.
Se adjunta el reporte de errores en las planillas procesadas.\n\n"""+ mensaje_error + "Saludos.")


    # Adjuntar los archivos
    for row in resultados_error:
        archivo_adjunto = os.path.join(carpeta_origen, row[2])
        with open(archivo_adjunto, "rb") as adjunto:
            mensaje.add_attachment(adjunto.read(), 
                                maintype="application", 
                                subtype="octet-stream", 
                                filename=row[2])
    
    # Enviar el correo
    contexto = ssl.create_default_context()
    with smtplib.SMTP_SSL(smtp_server, smtp_port, context=contexto) as server:
        server.login(email_origen, email_password)
        server.send_message(mensaje)
    
    print("Correo enviado 📧✅")


    # Mover los txt a la carpeta de errores
    for row in resultados_error:
        shutil.move(os.path.join(carpeta_origen, row[2]), carpeta_destino_error)